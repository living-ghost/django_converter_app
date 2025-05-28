import os
import subprocess
import uuid
import tempfile
from docx import Document
from django.contrib import messages
from django.conf import settings
from django.shortcuts import render, redirect
from django.contrib.auth import login, logout, authenticate
from .models import (
    User, DocxToPdfConversion, PdfToDocxConversion, 
    DocxToTxtConversion, JpgToPngConversion,
    PngToJpgConversion, ImgToPdfConversion )
from .utils import generate_otp, send_otp_email, send_reset_email, pwd_reset_success
from django.contrib.auth.decorators import login_required
from django.views.decorators.cache import never_cache
from django.http import FileResponse, HttpResponseBadRequest
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from pdf2docx import parse
from PyPDF2 import PdfReader
from PIL import Image
from itertools import chain
from django.utils.timezone import now


'''

    ########################################
    User Registration/OTP/Login Code Started
    ########################################

                                            '''


def before_login(request):
    if request.method == 'POST':
        if 'login_email' in request.POST:
            email = request.POST.get('login_email', '').strip()
            password = request.POST.get('login_password', '').strip()

            errors = {}

            if User.objects.filter(email=email).exists():
                user = authenticate(request, email=email, password=password)
                if user is not None:
                    login(request, user)
                    return redirect('after_login')
                else:
                    errors['login_error'] = 'Invalid email or password'
                    return render(request, "before_login.html", {
                        'errors': errors,  # Pass the errors dictionary
                        'show_login': True
                    })
            else:
                errors['no_user'] = 'Email not registered'
                return render(request, "before_login.html", {
                    'errors': errors,  # Pass the errors dictionary
                    'show_login': True
                })
            
        elif 'register_email' in request.POST:
            name = request.POST.get('register_name', '').strip()
            email = request.POST.get('register_email', '').strip()
            password = request.POST.get('register_password')
            confirm_password = request.POST.get('register_confirm_password')

            min_pass_length = 8

            # Validation
            errors = {}
            if not name:
                errors['name_error'] = 'Name is required'
            if not email:
                errors['email_error'] = 'Email is required'
            elif User.objects.filter(email=email).exists():
                errors['email_error'] = 'Email already exists'
            if not password:
                errors['password_error'] = 'Password is required'
            elif len(password) < min_pass_length:
                errors['password_error'] = f'Password must be at least {min_pass_length} characters long'
            if confirm_password != password:
                errors['confirm_error'] = 'Passwords do not match'

            if errors:
                return render(request, "before_login.html", {
                    **errors,
                    'show_register': True,
                    'show_login': False
                })
            
            # Generate otp and send OTP
            otp = generate_otp()
            send_otp_email(email, otp)

            # Store user input and OTP in session
            request.session['pending_user'] = {
                'name': name,
                'email': email,
                'password': password,
                'otp': str(otp)
            }

            return redirect('verify_otp')
        
        elif 'reset_email' in request.POST:
            user_email = request.POST.get('reset_email')
            if User.objects.filter(email=user_email).exists():
                send_reset_email(user_email)
                return redirect('before_login')
            else:
                print("User does not exist")


    # GET request or no form submitted
    return render(request, "before_login.html", {'show_register': False})

def verify_otp(request):
    if request.method == 'POST':
        if 'pending_user' in request.session:
            entered_otp = request.POST.get('otp')
            session_data = request.session.get('pending_user')

            if session_data and entered_otp == session_data['otp']:
                try:
                    user = User.objects.create_user(
                        email=session_data['email'],
                        full_name=session_data['name'],
                        password=session_data['password']
                    )
                    user = authenticate(request, email=session_data['email'], password=session_data['password'])
                    if user:
                        login(request, user)
                    del request.session['pending_user']
                    return redirect('after_login')
                except Exception:
                    return render(request, "otp/verify_otp.html", {
                        'error': 'Failed to create user. Try again.'
                    })

            return render(request, "otp/verify_otp.html", {
                'error': 'Invalid OTP'
            })
        
        elif 'reset_pwd' in request.session:
            entered_otp = request.POST.get('otp')
            session_data = request.session.get('reset_pwd')

            if session_data and entered_otp == session_data['otp']:
                try:
                    user = User.objects.get(email=session_data['email'])
                    user.set_password(session_data['new_password'])
                    user.save()
                    pwd_reset_success(session_data['email'])
                    del request.session['reset_pwd']
                    return redirect('before_login')
                except Exception:
                    return render(request, "otp/verify_otp.html", {
                        'error': 'Failed to update password. Try again.'
                    })

            return render(request, "otp/verify_otp.html", {
                'error': 'Invalid OTP'
            })

    return render(request, "otp/verify_otp.html")

@login_required
@never_cache
def after_login(request):
    user = request.user
    recent_docx_to_pdf = DocxToPdfConversion.objects.filter(user=user).exclude(pdf_file='').order_by('-uploaded_at')[:5]
    recent_pdf_to_docx = PdfToDocxConversion.objects.filter(user=user).exclude(docx_file='').order_by('-uploaded_at')[:5]
    recent_docx_to_txt = DocxToTxtConversion.objects.filter(user=user).exclude(txt_file='').order_by('-uploaded_at')[:5]
    recent_jpg_to_png = JpgToPngConversion.objects.filter(user=user).exclude(png_file='').order_by('-uploaded_at')[:5]
    recent_png_to_jpg = PngToJpgConversion.objects.filter(user=user).exclude(jpg_file='').order_by('-uploaded_at')[:5]
    recent_img_to_pdf = ImgToPdfConversion.objects.filter(user=user).exclude(pdf_file='').order_by('-uploaded_at')[:5]

    # Combine and sort all
    combined = sorted(
        chain(
            recent_docx_to_pdf,
            recent_pdf_to_docx,
            recent_docx_to_txt,
            recent_jpg_to_png,
            recent_png_to_jpg,
            recent_img_to_pdf,
        ),
        key=lambda x: x.uploaded_at,
        reverse=True
    )[:5]

    return render(request, 'after_login.html', {
        'recent_files': combined,
    })

@login_required
def user_logout(request):
    logout(request)
    return redirect('before_login')

def reset_password(request):
    if request.method == 'POST':
        email = request.POST.get('registered_email', '').strip()
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')

        if not User.objects.filter(email=email).exists():
            messages.error(request, "This email is not registered.")
            return render(request, "reset_password.html")

        if new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
            return render(request, "reset_password.html")

        # All checks passed
        otp = generate_otp()

        # Store user input and OTP in session
        request.session['reset_pwd'] = {
            'email': email,
            'new_'
            'password': new_password,
            'otp': str(otp)
        }

        send_otp_email(email, otp)

        return redirect('verify_otp')
    
    return render(request, "reset_password.html")


'''

    ####################################
    File Converter Coding Module Started
    ####################################

                                        '''


def docx_to_pdf(request):
    if request.method == 'POST' and request.FILES.get('file'):
        docx_file = request.FILES['file']

        # Check file extension
        if not docx_file.name.lower().endswith(('.docx', '.doc')):
            return HttpResponseBadRequest("Only .docx and .doc files are supported.")

        # Create database record first
        conversion = DocxToPdfConversion(user=request.user)
        
        # Save the original DOCX file to the database
        original_filename = docx_file.name
        conversion.docx_file.save(original_filename, ContentFile(docx_file.read()))
        
        # Reset file pointer after reading
        docx_file.seek(0)
        
        # Get the saved path from the database record
        docx_path = conversion.docx_file.path
        
        # Output directory for PDF
        output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', 'pdf')
        os.makedirs(output_dir, exist_ok=True)

        try:
            # Path to LibreOffice
            libreoffice_path = settings.LIBREOFFICE_PATH

            # LibreOffice command to convert
            command = [
                libreoffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ]
            
            # Run without shell=True for better security
            subprocess.run(command, check=True)

            # Construct PDF file path
            pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_filename)

            # Verify PDF was created
            if not os.path.exists(pdf_path):
                raise subprocess.CalledProcessError(1, command, "PDF file was not created")

            # Save PDF to database
            with open(pdf_path, 'rb') as pdf_file:
                conversion.pdf_file.save(pdf_filename, ContentFile(pdf_file.read()))
            
            # Serve the file as a download
            response = FileResponse(conversion.pdf_file.open('rb'), 
                                  as_attachment=True, 
                                  filename=pdf_filename)
            
            return response

        except subprocess.CalledProcessError as e:
            # Delete the database record if conversion fails
            conversion.delete()
            return HttpResponseBadRequest(f"Conversion failed: {str(e)}")
        except Exception as e:
            # Delete the database record if any other error occurs
            conversion.delete()
            return HttpResponseBadRequest(f"An error occurred: {str(e)}")

    return render(request, 'after_login.html')

def pdf_to_docx(request):
    if request.method == 'POST' and request.FILES.get('file'):
        pdf_file = request.FILES['file']

        if not pdf_file.name.lower().endswith('.pdf'):
            return HttpResponseBadRequest("Only PDF files are supported")

        with tempfile.TemporaryDirectory() as temp_dir:
            # Save uploaded PDF
            pdf_path = os.path.join(temp_dir, pdf_file.name)
            with open(pdf_path, 'wb+') as f:
                for chunk in pdf_file.chunks():
                    f.write(chunk)

            output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', 'docx')
            os.makedirs(output_dir, exist_ok=True)

            base_name = os.path.splitext(pdf_file.name)[0]
            docx_filename = f"{base_name}_converted.docx"  # Modified filename
            docx_path = os.path.join(output_dir, docx_filename)

            try:
                # 1. First convert with strict page dimensions
                parse(
                    pdf_path,
                    docx_path,
                    layout_engine='fixed',
                    kwargs={
                        'page_width': 8.27 * 72,  # A4 width in points (210mm)
                        'page_height': 11.69 * 72, # A4 height (297mm)
                        'margin': 0.5 * 72,       # 0.5 inch margins
                    },
                    text_settings={
                        'paragraph_space': 0,      # Remove extra spacing
                        'line_space': 1.0,         # Single line spacing
                        'orphan_control': True,    # Prevent lone lines
                    },
                    table_settings={
                        'split_strategy': 'avoid',  # Keep tables on one page
                    }
                )

                # 2. Post-process with python-docx to fix formatting
                from docx import Document
                from docx.shared import Pt, Inches
                
                doc = Document(docx_path)
                
                # Set document-wide page settings
                section = doc.sections[0]
                section.page_width = Inches(8.5)   # A4 width
                section.page_height = Inches(11.03)  # A4 height
                section.left_margin = Inches(0)
                section.right_margin = Inches(0)
                
                # Adjust paragraph formatting
                for paragraph in doc.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.keep_together = True  # Prevent splits
                    paragraph_format.keep_with_next = True  # Stay with next paragraph
                    paragraph_format.widow_control = True   # Prevent orphans
                
                doc.save(docx_path)

                # 3. Final validation
                if os.path.getsize(docx_path) < 1024:  # Check if file is too small
                    raise RuntimeError("Conversion failed - minimal output")

                # Save to database and return
                conversion = PdfToDocxConversion(user=request.user)
                with open(pdf_path, 'rb') as f:
                    conversion.pdf_file.save(pdf_file.name, ContentFile(f.read()))
                with open(docx_path, 'rb') as f:
                    conversion.docx_file.save(docx_filename, ContentFile(f.read()))

                return FileResponse(
                    open(docx_path, 'rb'),
                    as_attachment=True,
                    filename=docx_filename
                )

            except Exception as e:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                return HttpResponseBadRequest(f"Conversion error: {str(e)}")

    return render(request, 'after_login.html')

def docx_to_txt(request):
    if request.method == 'POST' and request.FILES.get('file'):
        docx_file = request.FILES['file']

        # Check file extension
        if not docx_file.name.lower().endswith(('.docx', '.doc')):
            return HttpResponseBadRequest("Only .docx and .doc files are supported.")

        # Create DB record first
        conversion = DocxToTxtConversion(user=request.user)
        original_filename = docx_file.name
        conversion.docx_file.save(original_filename, ContentFile(docx_file.read()))
        
        # Reset file pointer
        docx_file.seek(0)

        # Get saved path
        docx_path = conversion.docx_file.path

        # Output directory
        output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', 'txt')
        os.makedirs(output_dir, exist_ok=True)

        try:
            # Read DOCX file content
            document = Document(docx_path)
            text_content = '\n'.join([para.text for para in document.paragraphs])

            # Save as .txt
            txt_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.txt'
            txt_path = os.path.join(output_dir, txt_filename)
            with open(txt_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text_content)

            # Save to DB
            with open(txt_path, 'rb') as txt_file:
                conversion.txt_file.save(txt_filename, ContentFile(txt_file.read()))

            # Return as file download
            response = FileResponse(conversion.txt_file.open('rb'), 
                                    as_attachment=True, 
                                    filename=txt_filename)
            return response

        except Exception as e:
            conversion.delete()
            return HttpResponseBadRequest(f"An error occurred: {str(e)}")

    return render(request, 'after_login.html')

def jpg_to_png(request):
    if request.method == 'POST' and request.FILES.get('file'):
        jpg_file = request.FILES['file']

        # Check file extension
        if not jpg_file.name.lower().endswith('.jpg'):
            return HttpResponseBadRequest("Only .jpg files are supported.")

        # Create DB record
        conversion = JpgToPngConversion(user=request.user)
        original_filename = jpg_file.name
        conversion.jpg_file.save(original_filename, ContentFile(jpg_file.read()))

        # Reset pointer
        jpg_file.seek(0)

        # Saved path
        jpg_path = conversion.jpg_file.path

        # Output dir
        output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', 'png')
        os.makedirs(output_dir, exist_ok=True)

        try:
            # Convert using PIL
            with Image.open(jpg_path) as img:
                png_filename = os.path.splitext(os.path.basename(jpg_path))[0] + '.png'
                png_path = os.path.join(output_dir, png_filename)
                img.save(png_path, 'PNG')

            # Save converted file
            with open(png_path, 'rb') as png_file:
                conversion.png_file.save(png_filename, ContentFile(png_file.read()))

            # Serve file as download
            return FileResponse(conversion.png_file.open('rb'), 
                                as_attachment=True, 
                                filename=png_filename)

        except Exception as e:
            conversion.delete()
            return HttpResponseBadRequest(f"Conversion failed: {str(e)}")

    return render(request, 'after_login.html')

def png_to_jpg(request):
    if request.method == 'POST' and request.FILES.get('file'):
        png_file = request.FILES['file']

        # Check file extension
        if not png_file.name.lower().endswith('.png'):
            return HttpResponseBadRequest("Only .png files are supported.")

        # Create DB record
        conversion = PngToJpgConversion(user=request.user)  # You can rename this model to something generic if needed
        original_filename = png_file.name
        conversion.jpg_file.save(original_filename, ContentFile(png_file.read()))

        # Reset pointer
        png_file.seek(0)

        # Saved path
        png_path = conversion.jpg_file.path

        # Output dir
        output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', 'jpg')
        os.makedirs(output_dir, exist_ok=True)

        try:
            # Convert using PIL
            with Image.open(png_path) as img:
                # Convert to RGB to avoid transparency errors
                rgb_img = img.convert('RGB')
                jpg_filename = os.path.splitext(os.path.basename(png_path))[0] + '.jpg'
                jpg_path = os.path.join(output_dir, jpg_filename)
                rgb_img.save(jpg_path, 'JPEG')

            # Save converted file
            with open(jpg_path, 'rb') as jpg_file_obj:
                conversion.png_file.save(jpg_filename, ContentFile(jpg_file_obj.read()))

            # Serve file as download
            return FileResponse(conversion.png_file.open('rb'),
                                as_attachment=True,
                                filename=jpg_filename)

        except Exception as e:
            conversion.delete()
            return HttpResponseBadRequest(f"Conversion failed: {str(e)}")

    return render(request, 'after_login.html')

def img_to_pdf(request):
    if request.method == 'POST' and request.FILES.getlist('files'):
        uploaded_files = request.FILES.getlist('files')

        # Create a temporary directory
        temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_images')
        os.makedirs(temp_dir, exist_ok=True)

        image_objects = []
        conversion = ImgToPdfConversion(user=request.user)

        try:
            # Validate and save each file
            for uploaded_file in uploaded_files:
                if not uploaded_file.name.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
                    return HttpResponseBadRequest("Only image files (PNG, JPG, JPEG, BMP) are supported.")

                file_path = os.path.join(temp_dir, uploaded_file.name)
                with open(file_path, 'wb+') as destination:
                    for chunk in uploaded_file.chunks():
                        destination.write(chunk)

                # Open the image and convert to RGB
                image = Image.open(file_path).convert('RGB')
                image_objects.append(image)

            if not image_objects:
                return HttpResponseBadRequest("No valid images provided.")

            # Define PDF output path
            output_dir = os.path.join(settings.MEDIA_ROOT, 'converted', 'pdf')
            os.makedirs(output_dir, exist_ok=True)
            pdf_filename = 'converted_output.pdf'
            pdf_path = os.path.join(output_dir, pdf_filename)

            # Save all images as one PDF
            image_objects[0].save(pdf_path, save_all=True, append_images=image_objects[1:])

            # Save to model
            with open(pdf_path, 'rb') as pdf_file:
                conversion.pdf_file.save(pdf_filename, ContentFile(pdf_file.read()))
                conversion.save()

            # Return the file as response
            return FileResponse(conversion.pdf_file.open('rb'), as_attachment=True, filename=pdf_filename)

        except Exception as e:
            conversion.delete()
            return HttpResponseBadRequest(f"Conversion failed: {str(e)}")

    return render(request, 'after_login.html')